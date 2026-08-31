"""Builds docs/Solution_Walkthrough.docx — the competition submission
walkthrough.

Every figure in this document is traceable to a script in evaluation/ that
can be re-run. Where a result is weak, negative, or unstable, the document
says so; nothing here is rounded in our favour.

Run: PYTHONPATH=. .venv/bin/python docs/build_walkthrough.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x55, 0x5F, 0x6D)
WARN = RGBColor(0x8B, 0x2E, 0x2E)
OUT = Path("docs/Solution_Walkthrough.docx")


# --------------------------------------------------------------------------
# helpers
# --------------------------------------------------------------------------

def style_base(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.12
    for name, size, color, before in (
        ("Heading 1", 17, ACCENT, 20),
        ("Heading 2", 13, ACCENT, 14),
        ("Heading 3", 11.5, ACCENT, 10),
    ):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(4)


def para(doc, text, *, italic=False, size=None, color=None, after=6, align=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    return p


def rich(doc, parts, *, after=6):
    """parts: list of (text, bold, italic) — for inline emphasis."""
    p = doc.add_paragraph()
    for text, b, i in parts:
        r = p.add_run(text)
        r.bold = b
        r.italic = i
    p.paragraph_format.space_after = Pt(after)
    return p


def bullet(doc, text, *, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(3)
    return p


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def table(doc, headers, rows, *, widths=None, note=None, highlight_rows=()):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, "1F3A5F")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
            if ri in highlight_rows:
                r.bold = True
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
        if ri in highlight_rows:
            for c in cells:
                shade(c, "EEF3F9")
        elif ri % 2 == 1:
            for c in cells:
                shade(c, "F6F7F9")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    if note:
        para(doc, note, italic=True, size=8.5, color=MUTED, after=10)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def callout(doc, title, body, *, color=ACCENT):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    c.text = ""
    p1 = c.paragraphs[0]
    r = p1.add_run(title)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = color
    p2 = c.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    p2.paragraph_format.space_after = Pt(2)
    shade(c, "F4F6FA")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)
    return p


# --------------------------------------------------------------------------
# document
# --------------------------------------------------------------------------

def build() -> None:
    doc = Document()
    style_base(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(0.9)

    # ---- title -----------------------------------------------------------
    para(doc, "Solution Walkthrough", size=26, color=ACCENT, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Adaptive Red-Team / Blue-Team System for Agentic-Commerce Payment Fraud",
         size=13, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Mastercard Innovation Challenge 2026", size=10.5, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

    callout(
        doc, "Summary",
        "When an AI agent holds payment credentials and decides what to buy on a user's behalf, "
        "cryptography still protects execution — but nothing protects the decision that built the "
        "transaction. We built an end-to-end Identify → Generate → Defend system on a faithful "
        "simulation of Google's AP2 protocol: five attack families, an adaptive Red↔Blue loop where "
        "each side learns from the other, and a hybrid detector evaluated under a three-tier "
        "generalization protocol. Every figure below is reproducible from the repository, and the "
        "negative results are reported alongside the positive ones.")

    doc.add_page_break()

    # =====================================================================
    para(doc, "1.  Novel Fraud Attacks Identified", size=17, color=ACCENT, bold=True, after=6)

    para(doc,
         "Agentic commerce moves the attack surface. In a conventional payment flow the adversary must "
         "defeat authentication or cryptography. When an autonomous agent selects the product, requests "
         "the credentials, and constructs the mandate, an attacker who never touches a key can still "
         "change the outcome — because the mandate is signed faithfully around a decision that was "
         "already manipulated. Every attack below produces a cryptographically valid transaction.")

    para(doc, "1.1  The five implemented families", size=13, color=ACCENT, bold=True, after=4)

    table(doc,
          ["Family", "Mechanism", "Attacker capability", "Objective"],
          [
              ["Branded Whisper", "Hidden ranking directive embedded in product metadata",
               "Controls product descriptions only", "Payment manipulation"],
              ["Vault Whisper", "Social-engineered request for another user's stored credentials",
               "Is an authenticated user who can type", "Data exposure"],
              ["Intent Manipulation", "Plausible-but-worse decoy listing — no injected instruction anywhere",
               "Can list an ordinary product", "Payment manipulation"],
              ["Delegation Abuse", "Transaction executed outside its delegated scope (6 violation types)",
               "Controls a delegated sub-agent", "Payment manipulation"],
              ["Sequence Anomaly", "Compromised session: ATO burst, category drift, or slow drain",
               "Controls a compromised session", "Payment manipulation"],
          ],
          widths=[1.3, 2.7, 1.7, 1.3])

    para(doc, "1.2  What is genuinely novel here", size=13, color=ACCENT, bold=True, after=4)

    rich(doc, [("Automation of hand-crafted attacks. ", True, False),
               ("The two papers grounding this work hand-crafted one or two attacks per class and "
                "measured them over roughly ten trials. Both name the absence of automated generation "
                "as their own limitation. Our Red team generates, mutates and evolves those same "
                "attack classes automatically against a Blue detector that retrains in response.", False, False)])

    rich(doc, [("Intent manipulation without injection. ", True, False),
               ("A distinct family in which no adversarial instruction exists anywhere in the trace — "
                "every content item has contains_injection = False. The attack is purely catalogue "
                "ambiguity. This matters because every injection-detection defence is structurally "
                "blind to it.", False, False)])

    rich(doc, [("A quantified detectability frontier. ", True, False),
               ("We did not merely observe that slow-drain attacks evade detection; we established "
                "the point at which they become statistically undetectable by ANY method given the "
                "observable history, and validated that limit against closed-form statistical power "
                "analysis (Section 3.5).", False, False)])

    rich(doc, [("A demonstrated reward-hacking failure mode. ", True, False),
               ("Given a slightly wider search space, a reward-maximising optimiser abandoned stealth "
                "entirely and converged on high-volume ordinary spending — scoring 5.6× a genuine "
                "attack while being less harmful. We report this because it is a real hazard for any "
                "automated red-teaming system, and we hardened the reward against it.", False, False)])

    doc.add_page_break()

    # =====================================================================
    para(doc, "2.  How the System Generates and Simulates Attacks", size=17, color=ACCENT, bold=True, after=6)

    para(doc, "2.1  The simulated environment", size=13, color=ACCENT, bold=True, after=4)
    para(doc,
         "We model AP2's actual architecture rather than a generic marketplace: a Shopping Agent, "
         "Merchant Agent, Credentials Provider Agent and Merchant Payment Processor Agent, connected by "
         "a three-mandate chain (Intent → Cart → Payment). Two design choices are deliberate:")
    bullet(doc, "The environment is undefended by default. Mandates always sign; the credentials "
                "provider releases whatever it is asked for. Blue is the only thing between an attack "
                "and a completed transaction, which is what makes its contribution measurable.")
    bullet(doc, "A real LLM plays the Shopping Agent. Product selection and credential requests are "
                "genuine model calls, so agent decisions are actually manipulable rather than scripted "
                "to fail — several of our findings depend on the agent sometimes resisting.")

    para(doc, "2.2  The Red generation pipeline", size=13, color=ACCENT, bold=True, after=4)
    code(doc,
         "seed()      27 scenarios across 5 families — spending profiles, price tiers, categories,\n"
         "            cadences. Seed 0 of each family is paper-exact and frozen.\n"
         "   ↓\n"
         "mutate()    LLM families rewrite adversarial text; deterministic families move numeric\n"
         "            levers within clamped bounds. Driven by structured feedback: the specific\n"
         "            detection reasons from a caught round steer which lever moves next.\n"
         "   ↓\n"
         "simulate()  Run through the real AP2 chain against the live LLM agent.\n"
         "   ↓\n"
         "score()     R = intent_deviation x log1p(payment_impact) x realism x novelty\n"
         "                                                        - detection_probability\n"
         "   ↓\n"
         "memory      Deduplicate, measure novelty against everything already tried, retain\n"
         "            survivors as parents for the next generation.")

    para(doc,
         "The reward is multiplicative by design: an attack with no impact that evades perfectly still "
         "scores near zero, which blocks the degenerate \"never attack, never get caught\" strategy. Two "
         "terms required hardening after we observed them being exploited:")

    table(doc,
          ["Term", "Problem found", "Fix"],
          [
              ["payment_impact", "Grew linearly with transaction count, so emitting more transactions "
                                 "beat being stealthier",
               "log1p saturation — ordering preserved, marginal value of each added transaction falls"],
              ["realism", "Hardcoded constant, so it could not penalise anything despite being a "
                          "multiplicative term",
               "Measured per family against a benign profile (volume, cadence, amount plausibility)"],
          ],
          widths=[1.2, 2.8, 3.0])

    para(doc, "2.3  The adaptive Red ↔ Blue loop", size=13, color=ACCENT, bold=True, after=4)
    code(doc,
         "Generation N\n"
         "   ↓  Red population search against the CURRENT Blue\n"
         "   ↓  Pre-fit check — which attacks got through?\n"
         "   ↓  Split by LINEAGE ROOT into train / test\n"
         "   ↓  Blue fits (supervised) and calibrates (one-class) on TRAIN only\n"
         "   ↓  Post-fit recovery — did Blue close the gap it just had?\n"
         "   ↓  Metrics computed on TEST only\n"
         "Generation N+1 — Red mutates from the surviving lineages")

    para(doc, "Two disciplines make the resulting numbers trustworthy:")
    bullet(doc, "Lineage-root splitting. A mutated child is near-identical to its parent (we measured "
                "a 92% duplicate rate for one preset), so a naive per-trace split places near-twins on "
                "both sides. We hash the lineage root, keeping whole families of variants together.")
    bullet(doc, "Separated memories. Red's store suppresses similarity, which is what makes its novelty "
                "metric meaningful; Blue's store accumulates it, because training wants coverage. "
                "Sharing one store was a defect that silently starved Blue's training pool.")

    doc.add_page_break()

    # =====================================================================
    para(doc, "3.  Detection & Mitigation Model, with Efficacy Results", size=17, color=ACCENT, bold=True, after=6)

    para(doc, "3.1  Why a hybrid detector", size=13, color=ACCENT, bold=True, after=4)
    para(doc,
         "Our central technical finding is that supervised learning cannot generalise to an unseen "
         "attack strategy here — and this is provable rather than incidental. In a training pool "
         "containing only loud attacks, the feature that identifies a slow-drain attack carries no "
         "label information: benign values span and exceed the attack values. The coefficient on that "
         "feature is therefore statistically unidentified, and its fitted sign is set by sampling "
         "noise. No supervised model of any complexity can recover a coefficient the data does not "
         "constrain.")

    callout(doc, "How we caught this",
            "The same held-out experiment returned 0/13 in one run and 13/15 in another. Rather than "
            "report the favourable number, we traced the discrepancy to training-pool composition and "
            "established that we had been reading a coefficient-sign lottery as a result.",
            color=WARN)

    para(doc,
         "Blue therefore has two halves. The supervised half (logistic regression over six two-sided "
         "behavioural features) is strong on strategies it has labels for. The one-class half calibrates "
         "thresholds on attack-free traces only — it never sees an attack label, so a strategy it has "
         "never encountered is scored on exactly the same footing as a familiar one.")

    para(doc, "3.2  Four outcomes, not two", size=13, color=ACCENT, bold=True, after=4)
    para(doc, "This distinction governs every number that follows:")
    table(doc,
          ["", "Blue caught it", "Blue missed it"],
          [
              ["Attack succeeded", "Case B — harm, but caught", "Case C — harm, undetected"],
              ["Attack failed", "Case A — caught, low stakes", "Case D — blind spot, no harm yet"],
          ],
          widths=[1.5, 2.6, 2.6],
          note="Case C is the metric that matters. An attack that fails on its own is not a defence "
               "success, and high recall means little if the misses concentrate in the attacks that worked.")

    para(doc, "3.3  Adaptive loop results (3 generations, 164 traces, 156 LLM calls)",
         size=13, color=ACCENT, bold=True, after=4)
    table(doc,
          ["Gen", "Family", "Red ASR", "Recall", "FPR", "Case C", "F1", "Red reward"],
          [
              ["1", "reasoning_attack", "0.08", "0.00", "0.00", "0", "0.00", "−0.03"],
              ["2", "reasoning_attack", "0.33", "0.00", "0.00", "0", "0.00", "−0.52"],
              ["3", "reasoning_attack", "0.17", "0.00", "0.00", "0", "0.00", "−0.51"],
              ["1", "intent_manipulation", "0.00", "0.33", "0.00", "0", "0.50", "0.00"],
              ["2", "intent_manipulation", "0.00", "0.67", "0.00", "0", "0.80", "−0.75"],
              ["3", "intent_manipulation", "0.00", "0.72", "0.00", "0", "0.84", "−0.70"],
              ["1–3", "sequence_anomaly", "1.00→0.83", "1.00", "0.00", "0", "1.00", "−0.91"],
              ["1–3", "delegation_abuse (control)", "1.00", "1.00", "0.00", "0", "1.00", "n/a"],
          ],
          widths=[0.5, 1.7, 0.9, 0.7, 0.6, 0.6, 0.6, 0.9],
          highlight_rows=(6, 7))

    para(doc, "Four caveats, because this table flatters more than it should:", bold=True, after=3)
    bullet(doc, "intent_manipulation's zero Case C is trivial — no attack ever succeeded, so there was "
                "nothing to miss.")
    bullet(doc, "delegation_abuse's flat 1.00 is true by construction: its verifier provably covers all "
                "six violation types. It validates the harness; it is not a defensive achievement.")
    bullet(doc, "reasoning_attack shows 0.00 recall in later generations on test pools of n = 2–3. "
                "Case C is zero because the successes landed in the training split, not because Blue "
                "caught them.")
    bullet(doc, "Test pools are small throughout (n = 1–19). Single-generation numbers are indicative only.")

    para(doc, "Red's mean reward turns sharply negative by generation 2–3 in every learnable family — "
              "the clearest single signal that Blue wins the in-distribution arms race.", after=10)

    para(doc, "3.4  Generalization: three tiers that must never be conflated",
         size=13, color=ACCENT, bold=True, after=4)
    table(doc,
          ["Tier", "Question", "Supervised", "One-class", "Hybrid"],
          [
              ["1. In-distribution", "Attacks like those trained on?", "0.95", "0.87", "0.95"],
              ["2. Strength extrapolation", "Unseen parameters, known strategy?", "0.41", "0.04", "0.41"],
              ["3. Cross-strategy @0.85", "An entirely unseen strategy?", "0.05", "0.25", "0.25"],
              ["3. Cross-strategy @0.90", "(weaker attack)", "0.00", "0.15", "0.15"],
              ["Unseen loud strategies", "Sanity check on the one-class half", "—", "1.00", "1.00"],
          ],
          widths=[1.7, 2.3, 0.9, 0.9, 0.8],
          highlight_rows=(3, 4),
          note="Recall values. The one-class half reaches 1.00 on two attack strategies it was never "
               "trained on, having seen no attack labels at all — while the supervised half scores "
               "literally 0.00 on an unseen quiet strategy. Neither half dominates; the hybrid takes "
               "the better of the two at every tier without regression.")

    para(doc, "3.5  The slow-drain blind spot is an information limit, not a detector limit",
         size=13, color=ACCENT, bold=True, after=4)
    para(doc,
         "Rather than tune until the number improved, we asked whether the attack is detectable at all. "
         "Closed-form statistical power analysis and measurement agree closely:")
    table(doc,
          ["Attack strength", "Baseline history", "Predicted power", "Measured"],
          [
              ["0.90", "8 transactions (current)", "18%", "16%"],
              ["0.90", "30 transactions", "76%", "78%"],
              ["0.90", "60 transactions", "~99%", "98%"],
              ["0.95", "8 transactions (current)", "4%", "2%"],
          ],
          widths=[1.4, 2.0, 1.4, 1.4],
          note="With only 8 baseline transactions, an attack at multiplier ≤ 0.95 sits below the "
               "detection floor for any statistic whatsoever. This is the signal-to-noise ratio of the "
               "observable history, not a modelling failure — and it points at a concrete remedy "
               "(longer history) rather than a better classifier.")

    para(doc, "3.6  False-positive control", size=13, color=ACCENT, bold=True, after=4)
    para(doc,
         "A detector is only as good as its false-positive cost. Our null control (an attack-free "
         "sequence statistically identical to benign) initially reported a 20% false-positive rate. "
         "Investigation found two real causes and one that was not real at all:")
    bullet(doc, "The one-class layer estimated a Bonferroni-corrected 99th percentile from six samples — "
                "which simply returns the sample maximum. Fixed with a robust small-sample estimator.")
    bullet(doc, "The classifier used the library default 0.5 threshold while genuinely benign traces "
                "scored up to 0.89. Replaced with an explicit false-positive budget.")
    bullet(doc, "The measurement itself used n = 15, whose 95% confidence interval is [1%, 30%] — it "
                "could not distinguish 5% from 20%. It was reporting sampling noise.")

    table(doc,
          ["Amount multiplier", "Recall", "95% CI"],
          [
              ["0.85", "0.54", "[0.47, 0.60]"],
              ["0.90", "0.27", "[0.21, 0.33]"],
              ["0.95", "0.10", "[0.06, 0.14]"],
              ["1.00 (null control)", "FPR 0.04", "[0.02, 0.08]"],
          ],
          widths=[2.0, 1.4, 1.6],
          highlight_rows=(3,),
          note="n = 200 per point, target FPR 5%. Tradeoff stated plainly: fixing this cost recall at "
               "0.85 strength (~72% → ~54%). That is not a regression — the earlier recall was partly "
               "bought with false positives nobody was counting. The operating point is now an explicit, "
               "tunable parameter rather than a library default.")

    para(doc, "3.7  Mitigation: pre-execution enforcement", size=13, color=ACCENT, bold=True, after=4)
    para(doc,
         "Detection alone cannot claim prevented harm. We added a gate that runs before the sensitive "
         "action completes, keeping enforcement architecturally separate from detection so detectors "
         "remain reusable as pure classifiers:")
    table(doc,
          ["Control point", "Check", "Result"],
          [
              ["Credentials Provider", "Requested identity vs. authenticated session, before release",
               "Eliminates all cross-user exposure; legitimate same-identity requests still allowed"],
              ["Payment authorisation", "Six-clause delegation policy, before the transaction clears",
               "Attack success rate falls from 100% to 17%"],
          ],
          widths=[1.5, 2.6, 2.9],
          note="Traces record both the outcome and the counterfactual, so \"prevented by Blue\" is "
               "distinguishable from \"the attack simply failed\". The gate is implemented and tested "
               "but is not yet the default in the adaptive loop, which deliberately runs the "
               "undefended baseline so Blue's contribution stays measurable.")

    doc.add_page_break()

    # =====================================================================
    para(doc, "4.  Real-World Feasibility in Live Payment Environments", size=17, color=ACCENT, bold=True, after=6)

    para(doc, "4.1  What transfers directly", size=13, color=ACCENT, bold=True, after=4)
    table(doc,
          ["Component", "Why it is deployable", "Dependency"],
          [
              ["Caller-identity binding",
               "Compares the authenticated session against the identity whose data is requested, before "
               "release. Deterministic, explainable, immune to phrasing. Implements the PCAT P2 pattern "
               "from the protocol-attacks literature.",
               "Both identifiers already exist at the decision point in any real credential service"],
              ["Delegation policy verifier",
               "Six independent clauses (identity, scope, purpose, time, amount, chain). No training, no "
               "thresholds, fully auditable — the profile a payment network needs for a decline reason.",
               "A machine-readable delegation grant, which agentic-commerce protocols already define"],
              ["Two-sided CUSUM on spend",
               "A standard control-chart statistic, well understood in process monitoring, normalised "
               "against each account's own baseline rather than a population average.",
               "Transaction history of adequate length (see 4.2)"],
              ["Explicit FPR operating point",
               "Thresholds calibrated against a false-positive budget on attack-free traffic, which is "
               "how fraud systems are actually operated and reviewed.",
               "A representative sample of genuine traffic for calibration"],
          ],
          widths=[1.5, 3.1, 2.4])

    para(doc, "4.2  Deployment constraints our results imply", size=13, color=ACCENT, bold=True, after=4)
    bullet(doc, "Baseline history is the binding constraint, not model sophistication. Section 3.5 shows "
                "slow-drain detection moving from ~18% to ~98% purely by extending observable history "
                "from 8 to 60 transactions. Real institutions hold months of history — an advantage the "
                "simulator does not model and which likely makes production detection easier than our "
                "figures suggest.")
    bullet(doc, "The false-positive budget must be set deliberately. Our measured ROC is steep in the "
                "operating region: halving the false-positive rate roughly halves recall. This is a "
                "business decision about decline costs, not a modelling choice, and the system exposes "
                "it as an explicit parameter.")
    bullet(doc, "Baselines require managed re-establishment. A frozen per-account baseline cannot track "
                "legitimate long-term change in a customer's spending. Production deployment needs a "
                "verified re-baselining process, or genuine life changes will present as slow drains.")
    bullet(doc, "Unsupervised deviation is not the same as fraud. The one-class layer flags the unusual, "
                "not the malicious. It is appropriate for triage and review queues; it should not "
                "auto-decline without a second, higher-precision signal.")

    para(doc, "4.3  Honest limitations", size=13, color=ACCENT, bold=True, after=4)
    table(doc,
          ["Limitation", "Consequence for interpreting our results"],
          [
              ["Cross-strategy generalization to quiet attacks remains unsolved",
               "The one-class layer reaches 1.00 on unseen loud strategies but only 0.15–0.25 on unseen "
               "slow drains. Section 3.5 indicates the ceiling is environmental, not architectural."],
              ["A single victim model (gemini-3.1-flash-lite)",
               "Susceptibility is strongly model-dependent: we measured 0% attack success on ranking "
               "injection and 100% on identity override. These are not properties of \"LLM agents\" "
               "in general."],
              ["Small sample sizes",
               "A three-generation run yields tens of traces per family. The controlled evaluation "
               "suites are more trustworthy than any single live run, and we treat them as authoritative."],
              ["Simplified transaction schema",
               "No device fingerprint, geolocation, merchant reputation or session signal. Entire "
               "families of real-world detection are unavailable by construction — a production system "
               "would have strictly more to work with."],
              ["One detector remains keyword-based",
               "Branded Whisper detection is blind on roughly 90% of real attacks. We audited three "
               "candidate structural signals and rejected all three as either disguised keyword lists "
               "or properties the attacker directly controls, rather than ship a detector that "
               "overfits our own examples."],
          ],
          widths=[2.3, 4.7])

    para(doc, "4.4  Path to production", size=13, color=ACCENT, bold=True, after=4)
    table(doc,
          ["Priority", "Step", "Rationale"],
          [
              ["1", "Extend baseline history to 30–60 transactions",
               "Largest measured gain available; a data-availability change, not a modelling one"],
              ["2", "Make pre-execution enforcement the default path",
               "Converts classification into prevention with a measurable prevented-harm metric"],
              ["3", "Validate against a second victim model",
               "Establishes which findings are model-specific and which are structural"],
              ["4", "Integrate real transaction context (device, geo, merchant)",
               "Unlocks detection families the current schema cannot express"],
              ["5", "Add managed re-baselining",
               "Required before any per-account baseline can survive real customer behaviour"],
          ],
          widths=[0.7, 2.7, 3.6])

    callout(doc, "Reproducibility",
            "Every figure in this document is produced by a script in evaluation/ and can be re-run "
            "from the repository. The generalization suite and feature-validation harnesses require no "
            "API calls at all; the full adaptive loop costs approximately 156. Each script reports its "
            "own LLM call count, and the deterministic experiments assert that the count is zero.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
